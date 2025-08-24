"""
DOCX Document Handler for Healthcare Document Processing

Processes Microsoft Word documents with text extraction, metadata parsing,
table extraction, and healthcare-specific content analysis while maintaining HIPAA compliance.
"""

from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .base_handler import BaseDocumentHandler, DocumentMetadata, DocumentProcessingError


class DOCXDocumentHandler(BaseDocumentHandler):
    """
    Handles Microsoft Word DOCX document processing for healthcare applications

    Provides text extraction, metadata parsing, table extraction, and structured
    content analysis while maintaining healthcare compliance and PHI detection capabilities.
    """

    def __init__(self, enable_phi_detection: bool = True, enable_redaction: bool = False):
        """
        Initialize DOCX document handler

        Args:
            enable_phi_detection: Whether to perform PHI detection on content
            enable_redaction: Whether to create redacted versions of content
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

        super().__init__(enable_phi_detection, enable_redaction)

        # DOCX-specific configuration
        self.extract_tables = True
        self.extract_headers_footers = True
        self.extract_comments = False  # Comments may contain sensitive reviewer information
        self.preserve_formatting = True

    async def can_handle(self, file_path: str | Path, mime_type: str | None = None) -> bool:
        """
        Check if this handler can process the given DOCX document

        Args:
            file_path: Path to the document
            mime_type: MIME type of the document (optional)

        Returns:
            True if this is a DOCX document that can be processed
        """
        file_path = Path(file_path)

        # Check file extension
        if file_path.suffix.lower() in [".docx", ".docm"]:
            return True

        # Check MIME type
        if mime_type and mime_type in self.get_supported_mime_types():
            return True

        # Try to open as DOCX to verify format
        try:
            Document(file_path)
            return True
        except Exception:
            return False

    async def extract_content(self, file_path: str | Path) -> str:
        """
        Extract text content from DOCX document

        Args:
            file_path: Path to the DOCX document

        Returns:
            Extracted text content including paragraphs, tables, and headers/footers

        Raises:
            DocumentProcessingError: If DOCX processing fails
        """
        file_path = Path(file_path)

        try:
            doc = Document(file_path)
            extracted_content = []

            # Extract main document content
            main_content = await self._extract_document_body(doc)
            if main_content:
                extracted_content.append(main_content)

            # Extract tables if enabled
            if self.extract_tables:
                table_content = await self._extract_tables(doc)
                if table_content:
                    extracted_content.append(f"\n--- Tables ---\n{table_content}")

            # Extract headers and footers if enabled
            if self.extract_headers_footers:
                headers_footers = await self._extract_headers_footers(doc)
                if headers_footers:
                    extracted_content.append(f"\n--- Headers/Footers ---\n{headers_footers}")

            final_text = "\n".join(extracted_content).strip()

            if not final_text:
                self.logger.warning(f"No text extracted from DOCX: {file_path}")
                return ""

            return final_text

        except Exception as e:
            msg = f"Failed to extract content from DOCX: {e}"
            raise DocumentProcessingError(
                msg,
                str(file_path),
                e,
            )

    async def extract_metadata(self, file_path: str | Path) -> DocumentMetadata:
        """
        Extract metadata from DOCX document

        Args:
            file_path: Path to the DOCX document

        Returns:
            Document metadata including DOCX-specific properties
        """
        file_path = Path(file_path)

        try:
            file_stats = file_path.stat()
            doc = Document(file_path)

            # Extract DOCX core properties
            core_props = doc.core_properties

            # Build custom properties from DOCX metadata
            custom_properties = {
                "document_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
            }

            # Add core properties if available
            if core_props:
                custom_properties.update({
                    "title": getattr(core_props, "title", "") or "",
                    "author": getattr(core_props, "author", "") or "",
                    "subject": getattr(core_props, "subject", "") or "",
                    "keywords": getattr(core_props, "keywords", "") or "",
                    "comments": getattr(core_props, "comments", "") or "",
                    "category": getattr(core_props, "category", "") or "",
                    "created": str(getattr(core_props, "created", "") or ""),
                    "modified": str(getattr(core_props, "modified", "") or ""),
                    "last_modified_by": getattr(core_props, "last_modified_by", "") or "",
                    "version": getattr(core_props, "version", "") or "",
                    "revision": getattr(core_props, "revision", 0) or 0,
                })

            # Calculate content hash
            content_hash = await self._calculate_file_hash(file_path)

            return DocumentMetadata(
                file_name=file_path.name,
                file_size=file_stats.st_size,
                file_type="docx",
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content_hash=content_hash,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                last_modified=datetime.fromtimestamp(file_stats.st_mtime),
                custom_properties=custom_properties,
            )

        except Exception as e:
            msg = f"Failed to extract metadata from DOCX: {e}"
            raise DocumentProcessingError(
                msg,
                str(file_path),
                e,
            )

    async def _extract_document_body(self, doc: Document) -> str:
        """
        Extract text content from document body paragraphs

        Args:
            doc: python-docx Document instance

        Returns:
            Extracted text from document body
        """
        try:
            content = []

            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # This is a paragraph
                    para = Paragraph(element, doc)
                    text = para.text.strip()
                    if text:
                        content.append(text)
                elif isinstance(element, CT_Tbl) and not self.extract_tables:
                    # This is a table, but we're not extracting tables separately
                    table = Table(element, doc)
                    table_text = await self._extract_single_table(table)
                    if table_text:
                        content.append(f"[TABLE]\n{table_text}\n[/TABLE]")

            return "\n\n".join(content) if content else ""

        except Exception as e:
            self.logger.warning(f"Failed to extract document body: {e}")
            return ""

    async def _extract_tables(self, doc: Document) -> str:
        """
        Extract content from all tables in the document

        Args:
            doc: python-docx Document instance

        Returns:
            Formatted table content as text
        """
        try:
            table_contents = []

            for table_num, table in enumerate(doc.tables, 1):
                table_text = await self._extract_single_table(table)
                if table_text:
                    table_contents.append(f"Table {table_num}:\n{table_text}")

            return "\n\n".join(table_contents) if table_contents else ""

        except Exception as e:
            self.logger.warning(f"Failed to extract tables: {e}")
            return ""

    async def _extract_single_table(self, table: Table) -> str:
        """
        Extract content from a single table

        Args:
            table: python-docx Table instance

        Returns:
            Formatted table content as text
        """
        try:
            table_data = []

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # Replace internal newlines with spaces for table formatting
                    cell_text = cell_text.replace("\n", " ").replace("\r", " ")
                    row_data.append(cell_text)

                if any(cell.strip() for cell in row_data):  # Skip empty rows
                    table_data.append(" | ".join(row_data))

            return "\n".join(table_data) if table_data else ""

        except Exception as e:
            self.logger.warning(f"Failed to extract single table: {e}")
            return ""

    async def _extract_headers_footers(self, doc: Document) -> str:
        """
        Extract content from headers and footers

        Args:
            doc: python-docx Document instance

        Returns:
            Formatted headers and footers content as text
        """
        try:
            headers_footers_content = []

            for section_num, section in enumerate(doc.sections, 1):
                section_content = []

                # Extract header content
                header = section.header
                if header and header.paragraphs:
                    header_text = "\n".join([p.text.strip() for p in header.paragraphs if p.text.strip()])
                    if header_text:
                        section_content.append(f"Header: {header_text}")

                # Extract footer content
                footer = section.footer
                if footer and footer.paragraphs:
                    footer_text = "\n".join([p.text.strip() for p in footer.paragraphs if p.text.strip()])
                    if footer_text:
                        section_content.append(f"Footer: {footer_text}")

                if section_content:
                    headers_footers_content.append(f"Section {section_num}:\n" + "\n".join(section_content))

            return "\n\n".join(headers_footers_content) if headers_footers_content else ""

        except Exception as e:
            self.logger.warning(f"Failed to extract headers/footers: {e}")
            return ""

    def _get_content_type(self) -> str:
        """Get the content type identifier for DOCX handler"""
        return "docx_document"

    def get_supported_formats(self) -> list[str]:
        """Get list of supported DOCX file formats/extensions"""
        return [".docx", ".docm"]

    def get_supported_mime_types(self) -> list[str]:
        """Get list of supported MIME types for DOCX"""
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/vnd.ms-word.document.macroEnabled.12",
        ]

    async def extract_structured_content(self, file_path: str | Path) -> dict[str, Any]:
        """
        Extract structured content from DOCX including paragraphs, tables, and formatting

        Args:
            file_path: Path to the DOCX document

        Returns:
            Dictionary containing structured document content
        """
        file_path = Path(file_path)

        try:
            doc = Document(file_path)

            structured_content = {
                "paragraphs": [],
                "tables": [],
                "headers_footers": {},
                "document_properties": {},
            }

            # Extract paragraphs with basic formatting info
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_info = {
                        "paragraph_number": para_num + 1,
                        "text": paragraph.text.strip(),
                        "style": paragraph.style.name if paragraph.style else "Normal",
                        "alignment": str(paragraph.alignment) if paragraph.alignment else "Unknown",
                    }
                    structured_content["paragraphs"].append(para_info)

            # Extract tables with structure
            for table_num, table in enumerate(doc.tables):
                table_data = {
                    "table_number": table_num + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "content": [],
                }

                for row_num, row in enumerate(table.rows):
                    row_data = {
                        "row_number": row_num + 1,
                        "cells": [cell.text.strip() for cell in row.cells],
                    }
                    table_data["content"].append(row_data)

                structured_content["tables"].append(table_data)

            # Extract headers and footers by section
            for section_num, section in enumerate(doc.sections):
                section_key = f"section_{section_num + 1}"

                header_text = ""
                if section.header and section.header.paragraphs:
                    header_text = "\n".join([p.text.strip() for p in section.header.paragraphs if p.text.strip()])

                footer_text = ""
                if section.footer and section.footer.paragraphs:
                    footer_text = "\n".join([p.text.strip() for p in section.footer.paragraphs if p.text.strip()])

                structured_content["headers_footers"][section_key] = {
                    "header": header_text,
                    "footer": footer_text,
                }

            # Extract document properties
            if doc.core_properties:
                structured_content["document_properties"] = {
                    "title": getattr(doc.core_properties, "title", "") or "",
                    "author": getattr(doc.core_properties, "author", "") or "",
                    "subject": getattr(doc.core_properties, "subject", "") or "",
                    "created": str(getattr(doc.core_properties, "created", "") or ""),
                    "modified": str(getattr(doc.core_properties, "modified", "") or ""),
                }

            return structured_content

        except Exception as e:
            msg = f"Failed to extract structured content from DOCX: {e}"
            raise DocumentProcessingError(
                msg,
                str(file_path),
                e,
            )
